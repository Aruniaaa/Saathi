import fitz
import docx
from google import genai
from dotenv import load_dotenv
import os
from langchain_core.documents import Document
from langchain.agents import create_agent
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.tools import tool
from langchain_community.vectorstores import Chroma
from langchain_google_genai import ChatGoogleGenerativeAI



load_dotenv()

api_key = os.getenv("GEMINI_KEY")

embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004", google_api_key=api_key)


vectorstore = Chroma(
    collection_name="Quizzes",
    embedding_function=embeddings
)

llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash-lite", google_api_key=api_key)


load_dotenv()

api_key  = os.getenv("GEMINI_KEY")

client = genai.Client(api_key=api_key)



def process(query, context):
    try:
        

        prompted_query = f"""
        You are Saathi.
        You are a mentor specializing in STEM — science, technology, engineering, and mathematics.
        Your purpose is teaching and guiding, never directly solving. You are a sophisticated, smart “rubber duck”—a partner in reasoning and understanding.
    
        Rules:
    
        Explain, Don’t Answer
        - Always break down the *why* behind any concept, derivation, or problem.
        - Focus on the underlying laws, patterns, and logic — not just formulas or results.
        - Avoid directly solving or computing the final answer. Instead, help the user understand the logic
    
        Teach Theory
        - Always connect the question to fundamental scientific or mathematical principles.
        - Relate it to real-world intuition: why it works, where it fails, and how it connects to other ideas.
        - Prioritize explanations that build conceptual depth and long-term understanding.
    
        Guide, Don’t Solve
        - Never provide step-by-step solutions, code blocks, or final numeric results.
        - Instead, ask Socratic questions, hint toward relevant formulas or laws, and guide through conceptual checkpoints.
        - Encourage the learner to derive and reason, not memorize or copy.
    
        Encourage Cross-Disciplinary Thinking
        - Show how physics connects to math, how biology links with chemistry, or how computer science overlaps with logic.
        - Promote cognitive discipline — learning how to *think*, not what to memorize.
    
        Strict Prohibitions:
        - No giving final answers, code blocks, complete solutions, or formula substitutions.
        - No oversimplifying to the point of removing reasoning.
        - No workarounds that skip understanding.


        CONTEXT START

        {context}

        CONTEXT END

        Query: {query}
        """ 

        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompted_query
        )

        return response.text
    except Exception as e:
        print(e)




def read_pdf(file):
    
    doc = fitz.open(stream=file.read(), filetype="pdf")
    all_text = []
    for page in doc:
        
        text = page.get_text("text")
        text = " ".join(text.split()) 
        all_text.append(text)  


    return "\n\n".join(all_text).strip()

def read_docx(file):

    doc = docx.Document(file)
    all_text = []

    for page in doc.paragraphs:
        text = page.text
        if text != "":
            all_text.append(text)
    
    return "\n\n".join(all_text).strip()



def return_prompt(file=None, text=None):

    content = ""

    if file:

        ext = file.name.lower().split(".")[-1]

        if ext == "pdf":
            content = read_pdf(file)
        
        elif ext == "docx":
            content = read_docx(file)

        elif ext == "txt":
            content = file.read().decode("utf-8")

    if text:
        content += "\n" + text

    

    
    prompt = f"""You are generating a multiple-choice quiz strictly based on the provided content.

        Rules:
        - The number of questions must depend on the content length (minimum 5, maximum 20).
        - If the content is too short, produce exactly 5 questions.
        - Use ONLY information from the content. Do not add external knowledge.
        - Every question must be answerable using the content alone.
        - Follow the output format exactly with no extra text.

        Output format (strict):
        Return a JSON array.
        Each question must be a JSON object containing these fields:

        "number": integer

        "question": string

        "options": list of 4 strings in order [A, B, C, D]

        "answer": string ("A" | "B" | "C" | "D")

        Do not include explanations, comments, or summaries. Output raw JSON only.
        No commentary.
        No markdown.
        No backticks.
        The response must be a valid JSON array that can be parsed with json.loads().

        CONTENT START
        {content}
        CONTENT END
        """

    return prompt


def get_quiz(prompt):

    response = client.models.generate_content(
        model="gemini-2.5-flash-lite",
        contents=prompt,
        config={
        "response_mime_type": "application/json"
        }
    )

    return response.text



def store_in_vdb(wrong_questions, user_id):

    docs = [
        Document(
            page_content=wrong_questions,
            metadata={
                "user_id": user_id
            }
        )
    ]

    vectorstore.add_documents(docs)


@tool
def retrieve_similar_quizzes(user_id: str, wrong_questions: str):
    """Get similar quizzes to the last quiz that the user took to analyze patterns and generate feedback.

    Args:
        user_id: The unique id of the user whose patterns and habits we're analyzing.
        wrong_questions: A string containing what questions the user got wrong,
                         what the correct answer was, and what the user's answer was.

    Returns:
        A list of similar quiz documents from the vector database.
    """


    response = client.models.generate_content(
        model="gemini-2.5-flash-lite",
        contents=f"""Here is some info on a quiz, return a 3-5 word query for a RAG search to get similar quizzes so that
        an AI agent can give the user feedback on similar quizzes and analyse where they went wrong. Return words like 'algebra' or 'trignometry' or 'economics'.
        Output ONLY the query, no explanation, no punctuation.
        Here is the quiz info : {wrong_questions}"""
        )

    results = vectorstore.similarity_search(
        query=response.candidates[0].content.parts[0].text,
        k=3,
        filter={"user_id": user_id}
    )

    print(results)

    return results




agent = create_agent(
    model=llm,
    tools=[retrieve_similar_quizzes],
    system_prompt="You are Saathi, you analyze student quiz performance, retrieve similar past quizzes, detect patterns, and give personalized study advice. Use the retrieval tool when needed."
)






